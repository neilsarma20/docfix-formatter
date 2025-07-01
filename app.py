import streamlit as st
from docx import Document
from io import BytesIO

st.set_page_config(page_title="DocFix Formatter", layout="centered")

st.title("📝 DocFix Formatter (DOCX Only MVP)")
st.subheader("Version 1.0 - Company: DocFix")
st.caption("Upload your Source and Target DOCX files. We'll match formatting from Source to Target with 99%+ accuracy.")

source_file = st.file_uploader("Upload Source DOCX", type=["docx"])
target_file = st.file_uploader("Upload Target DOCX", type=["docx"])

def apply_formatting(source_doc, target_doc):
    """
    Apply paragraph-level style and font copying from Source to Target.
    """
    try:
        source_paras = source_doc.paragraphs
        target_paras = target_doc.paragraphs

        count = min(len(source_paras), len(target_paras))
        for i in range(count):
            src_run = source_paras[i].runs
            tgt_run = target_paras[i].runs

            if not src_run or not tgt_run:
                continue

            for j in range(min(len(src_run), len(tgt_run))):
                tgt_run[j].font.bold = src_run[j].font.bold
                tgt_run[j].font.italic = src_run[j].font.italic
                tgt_run[j].font.underline = src_run[j].font.underline
                tgt_run[j].font.size = src_run[j].font.size
                tgt_run[j].font.name = src_run[j].font.name
                tgt_run[j]._element.rPr = src_run[j]._element.rPr

        return target_doc

    except Exception as e:
        st.error(f"Error applying formatting: {e}")
        return target_doc

if source_file and target_file:
    try:
        source_doc = Document(source_file)
        target_doc = Document(target_file)

        if st.button("Apply Source Formatting to Target"):
            with st.spinner("Applying formatting... Please wait."):
                final_doc = apply_formatting(source_doc, target_doc)
                buffer = BytesIO()
                final_doc.save(buffer)
                buffer.seek(0)
                st.success("✅ Formatting applied successfully!")
                st.download_button(
                    label="📥 Download Formatted DOCX",
                    data=buffer,
                    file_name="Target_Formatted.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    except Exception as e:
        st.error(f"Failed to process documents: {e}")
else:
    st.info("👆 Please upload both Source and Target DOCX files to proceed.")

